#!/usr/bin/env python3
import docx
import re

def docx_to_markdown(docx_path, md_path):
    """将docx文件转换为markdown格式"""
    doc = docx.Document(docx_path)
    md_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_content.append('')
            continue
        
        # 处理标题
        if para.style and para.style.name.startswith('Heading'):
            level = int(re.search(r'Heading (\d+)', para.style.name).group(1))
            md_content.append('#' * level + ' ' + text)
        else:
            md_content.append(text)
    
    # 处理表格
    for table in doc.tables:
        table_content = []
        # 处理表头
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        table_content.append('| ' + ' | '.join(headers) + ' |')
        table_content.append('| ' + ' | '.join(['---' for _ in headers]) + ' |')
        # 处理表格内容
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            table_content.append('| ' + ' | '.join(cells) + ' |')
        md_content.extend(table_content)
        md_content.append('')
    
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_content))
    
    print(f'转换完成: {docx_path} -> {md_path}')

if __name__ == '__main__':
    import sys
    if len(sys.argv) != 3:
        print('用法: python convert_docx_to_md.py <input.docx> <output.md>')
        sys.exit(1)
    docx_path = sys.argv[1]
    md_path = sys.argv[2]
    docx_to_markdown(docx_path, md_path)